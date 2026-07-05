import traceback
import asyncio
import os

def parse_pdf(file_path: str) -> str:
    try:
        import pdfplumber
        text_parts = []
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    text_parts.append(text)
        return "\n\n".join(text_parts)
    except Exception as e:
        traceback.print_exc()
        return ""

def parse_docx(file_path: str) -> str:
    try:
        from docx import Document
        doc = Document(file_path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        if paragraphs:
            return "\n\n".join(paragraphs)
        # Fallback: extract text from tables
        table_texts = []
        for table in doc.tables:
            for row in table.rows:
                row_texts = [cell.text for cell in row.cells]
                table_texts.append(" | ".join(row_texts))
        return "\n".join(table_texts)
    except Exception as e:
        traceback.print_exc()
        return ""

async def parse_file(file_path: str) -> str:
    try:
        ext = os.path.splitext(file_path)[1].lower()
        if ext == ".pdf":
            return parse_pdf(file_path)
        elif ext == ".docx":
            return parse_docx(file_path)
        else:
            return ""
    except Exception as e:
        traceback.print_exc()
        return ""

if __name__ == "__main__":
    import sys
    if len(sys.argv) > 1:
        result = asyncio.run(parse_file(sys.argv[1]))
        print(repr(result))
    else:
        print("Usage: python -m tools.file_parser <file_path>")
